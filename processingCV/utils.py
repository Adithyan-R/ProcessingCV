# utils.py
import PyPDF2
import os
import docx  # for .docx files
from docx import Document


def extract_text_from_pdf(pdf_path):
    text = ''
    with open(pdf_path, 'rb') as pdf_file:
        reader = PyPDF2.PdfReader(pdf_file)
        for page_num in range(len(reader.pages)):
            text += reader.pages[page_num].extract_text()
    return text


def extract_text_from_docx(file_path):
    doc = docx.Document(file_path)
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    return text


#ANTIWORD_PATH = '/full/path/to/antiword'

    # Use antiword for extracting text from .doc files


def extract_text_from_doc(file_path):
    try:
        doc = Document(file_path)
        text = '\n'.join([para.text for para in doc.paragraphs])
        return text
    except Exception as e:
        print(f"Error extracting text from .docx file: {e}")
        return ''





'''
def extract_text_from_doc(file_path):
    try:
        process = subprocess.Popen(['antiword', file_path], stdout=subprocess.PIPE, stderr=subprocess.PIPE)
        stdout, stderr = process.communicate()
        if stderr:
            print("Error output from antiword:")
            print(stderr.decode('utf-8'))
        text = stdout.decode('utf-8')
        return text
    except Exception as e:
        print(f"Error extracting text from .doc file: {e}")
        return ''

def extract_text_from_doc(file_path):
    try:
        text = textract.process(file_path).decode('utf-8')
        return text
    except Exception as e:
        print(f"Error extracting text from .doc file: {e}")
        return ''







def extract_text_from_doc(doc_file):
    # Convert .doc file to text using antiword
    temp_text_file = os.path.splitext(doc_file.name)[0] + ".txt"
    subprocess.run(["antiword", doc_file.path, ">", temp_text_file], shell=True)
    # Read text from the generated text file
    with open(temp_text_file, 'r') as f:
        text = f.read()
    # Clean up temporary text file
    os.remove(temp_text_file)
    return text

'''